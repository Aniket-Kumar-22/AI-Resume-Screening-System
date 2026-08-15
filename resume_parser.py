import os 
from pathlib import Path
from dotenv import load_dotenv
from groq import Groq
import time
from pydantic import BaseModel

load_dotenv()
my_api_key = os.getenv("Qroq_Api_Key")

if not my_api_key:
    raise ValueError("api_error")

client = Groq(api_key = my_api_key)
model = "llama-3.3-70b-versatile"

job_description ="""
Our company needs to hire an experienced and talented Machine Learning and Artificial Intelligence (ML/AI) Engineer. This role will be crucial to the development and implementation of cutting-edge AI products. Its responsibilities will include designing and constructing sophisticated machine learning models, in addition to improving and updating existing systems. The ultimate goal is to create efficient self-learning applications that can evolve over time. If you want to be at the forefront of machine learning innovation, we hope you’ll join our team. 

Typical Duties and Responsibilities
Design machine learning systems
Research and implement machine learning algorithms and tools
Manage and direct research and development processes to meet the needs of our AI strategy
Develop machine learning applications in alignment with project requirements and business goals
Perform machine learning tests and statistical analysis in order to fine-tune the machine learning systems
Select appropriate datasets and data representation methods
Extend existing machine learning libraries and frameworks
Train systems and retrain as necessary
Work with the engineering and leadership teams on the functional design, process design, prototyping, testing, and training of AI/ML solutions
Advise leaders on technology, strategy, and policy issues related to AI/ML
Education
Bachelor’s degree in computer science, mathematics, or a related field
Required Skills and Experience
2+ years of experience applying AI to practical uses
Experience with deep learning, NLP, and TensorFlow
Experience writing robust code in Python, Java, and/or R
Experience in REST API development, NoSQL database design, and RDBMS design and optimizations
Knowledge of basic algorithms and object-oriented and functional design principles
Knowledge of data structures, data modeling, and software architecture
Knowledge of math, probability, statistics, and algorithms
Knowledge of machine learning frameworks such as Keras and PyTorch
Knowledge of machine learning libraries such as scikit-learn
Excellent communication skills
Strong analytical and problem solving skills
Preferred Qualifications
Master’s degree in a relevant technology field
Experience with cloud environments
"""
class JobD(BaseModel):
    role : str
    required_skills: list[str]
    preferred_skills:list[str]
    minimum_experience: float | None
    education_requirements: list[str]
    responsibilities: list[str]

jobd_schema = JobD.model_json_schema()

system_prompt = f"""
You are an expert HR assistant.

Your job is to analyze job description and extract structured information from them.

Return ONLY valid JSON matching this schema:{jobd_schema}

IMPORTANT:
Do NOT return the schema itself.
Do NOT return fields like "properties","title" or "type".
Fill the schema with actual information extracted from the job description.

If minimum experience is not mentioned, return null.
If information for a list is missing, return an empty list.
Do not invent information.
"""

user_prompt = f"""
Analyze the following job description:
{job_description}
"""

message_system = {
    "role" : "system",
    "content" : system_prompt
}

message_users = {
    "role" : "user",
    "content" : user_prompt
}
response_format ={
    "type" : "json_object"
}

messages = [message_system, message_users]

response = client.chat.completions.create(model=model, messages=messages, response_format=response_format)

answer = response.choices[0].message.content

raw_json = answer
#print(raw_json)

import json
job_data = json.loads(raw_json)

job = JobD(**job_data)

print(job.minimum_experience)
print(job.education_requirements)


# Parse real
class MatchResult(BaseModel):
    score : float
    details: dict

class Experience(BaseModel):
    company : str | None = None
    role : str | None = None
    duration : str | None = None
    description : str | None = None
    skills_used : list[str] = []

class Resume(BaseModel):
    name : str | None = None
    email : str | None = None 
    phone : str | None = None

    total_experience_years : float | None = None

    skills: list[str] = []
    experience: list[str] = []
    education: list[str] = []
    projects: list[str] = []
    certifications: list[str] = []

resume_schema = Resume.model_json_schema()
def final_score(job,resume):
    match_schema = MatchResult.model_json_schema()
    prompt = f"""
    You are an HR recruitor.

    Compare the candidate's resume with the job description.

    JOB DESCRIPTION:
    {job.model_dump_json(indent=2)}

    CANDIDATE RESUME:
    {resume.model_dump_json(indent=2)}
    Return JSON matching this schema:
    {match_schema}

    Give me:

    1.Candidate name
    2.Matching Skills
    3.Missing important skills
    4.Whether experience requirement is met
    5.Overall match percentage from 0 to 100
    6.A short final verdict

    Keep the response concise and easy to read.
    """

    message={
        "role": "user",
        "content": prompt
    }
    messages=[message]
    response_format={
        "type": "json_object"
    }
    response = client.chat.completions.create(model=model, messages=messages, response_format=response_format)
    data = json.loads(response.choices[0].message.content)
    return MatchResult(**data)

def parse_resume(resume_text):
    system_prompt = f"""
    You are an expert resume parser.

    Extract information from the resume based on its maeaning,
    not only based on exact section headings.

    Different resumes may use different headings.

    For example:
    - Experience 
    - Professional Experience
    - Work History
    - Employment
    - Internships

    These may all contain relevant experience.

    Skills may also appear in the skills section, work experience,
    internships or projects.

    Return only valid JSON matching this schema:
    {resume_schema}

    Important rules:
    1.Do not invent information.
    2.If a value is not available, return null.
    3.If a list has no information, return an empty list.
    4.Include internships inside experiences.
    5.Extract skills mentioned across the entire resume. 
    """
    user_prompt = f"""
    Parse the following resume:
    {resume_text}
    """
    message_system={
        "role" : "system",
        "content" : system_prompt
    }
    message_user={
        "role" : "user",
        "content" : user_prompt
    }
    messages=[message_system, message_user]
    response_format={
        "type": "json_object"
    }
    response=client.chat.completions.create(model=model, messages=messages, response_format=response_format)
    raw_output = response.choices[0].message.content
    data = json.loads(raw_output)
    resume = Resume(**data)
    return resume

from pypdf import PdfReader
from docx import Document
def read_pdf(file_path):
    reader = PdfReader(file_path)
    text =""
    for page in reader.pages:
        page_text = page.extract_text()
        if page_text:
            text += page_text + "\n"
    return text

def read_docx(file_path):
    document = Document(file_path)
    text = ""
    for paragraph in document.paragraphs:
        if paragraph.text.strip():
            text += paragraph.text + "\n"

    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text.strip():
                    text += cell.text + "\n"
    return text


def read_resume(file_path):
    if file_path.suffix.lower() == ".pdf":
        return read_pdf(file_path)
    elif file_path.suffix.lower() == ".docx":
        return read_docx(file_path)
    else:
        return None


resume_folder = Path("resumes")
all_results=[]
for file_path in resume_folder.iterdir():
    if file_path.suffix.lower() not in [".pdf",".docx"]:
        continue
    print("\nProcessing:", file_path.name)
    resume_text = read_resume(file_path)
    parsed_text = parse_resume(resume_text)
    time.sleep(5)
    result = final_score(job, parsed_text)
    time.sleep(5)
    print("Score:",result.score)
    all_results.append({
        "name": parsed_text.name,
        "score": result.score,
        "details": result.details
    })
all_results.sort(
    key=lambda candidate: candidate["score"],
    reverse=True
)
top_2 = all_results[:2]
worst_2= all_results[-2:]

print("TOP 2 CANDIDATES")
for candidate in top_2:
    print(
        candidate["name"],
        "-",
        candidate["score"],
        "%"
    )
    print(candidate["details"])

print("LOWEST 2 CANDIDATES")
for candidate in worst_2:
    print(
        candidate["name"],
        "-",
        candidate["score"],
        "%"
    )
    print(candidate["details"])


    """
    uv venv
    .\.venv\Scripts\Activate.ps1
    
    uv add python-dotenv
    uv add groq
    uv add pydantic
    uv add pypdf
    uv add python-docx
    """